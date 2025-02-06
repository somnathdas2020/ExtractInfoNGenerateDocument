import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

# Image source folder
IMAGE_FOLDER = r"C:\Users\somnath.das\Desktop\Default\Images"

def html_color_to_rgb(color_code):
    """Converts HTML color codes (#RRGGBB) to RGB format for docx."""
    if color_code and color_code.startswith("#") and len(color_code) == 7:
        return tuple(int(color_code[i:i+2], 16) for i in (1, 3, 5))
    return None

def add_hyperlink(paragraph, text, url):
    """Adds a clickable hyperlink in the Word document."""
    part = paragraph.part
    r_id = part.relate_to(url, "hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._element.append(hyperlink)
    return hyperlink

def extract_content_from_html(file_path, doc):
    """Extracts formatted content only from id='mainbody', preserving structure and colors."""
    with open(file_path, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file, "html.parser")

        # Extract only the content inside id="mainbody"
        mainbody = soup.find(id="mainbody")
        if not mainbody:
            print(f"Warning: 'mainbody' ID not found in {file_path}")
            return

        # Process each element inside mainbody
        for element in mainbody.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "b", "strong", "i", "em", "ul", "ol", "li", "a", "span", "img"], recursive=True):
            # Extract text color if available
            color = None
            if "style" in element.attrs:
                styles = element["style"].split(";")
                for style in styles:
                    if "color:" in style:
                        color_code = style.split(":")[-1].strip()
                        color = html_color_to_rgb(color_code)

            # Apply text formatting based on element type
            if element.name.startswith("h"):  # Headings
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text(), level=level)
                if color:
                    heading.runs[0].font.color.rgb = RGBColor(*color)
            elif element.name in ["b", "strong"]:  # Bold text
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.bold = True
                if color:
                    run.font.color.rgb = RGBColor(*color)
            elif element.name in ["i", "em"]:  # Italic text
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.italic = True
                if color:
                    run.font.color.rgb = RGBColor(*color)
            # elif element.name == "p":  # Paragraphs
            #     p = doc.add_paragraph(element.get_text())
            #     if color:
            #         p.runs[0].font.color.rgb = RGBColor(*color)
            elif element.name == "ul":  # Bulleted List (Fixed duplicate issue)
                for li in element.find_all("li", recursive=False):  # Only direct children
                    p = doc.add_paragraph(li.get_text(), style="ListBullet")
                    if color:
                        p.runs[0].font.color.rgb = RGBColor(*color)
            elif element.name == "ol":  # Numbered List
                for li in element.find_all("li", recursive=False):  # Only direct children
                    p = doc.add_paragraph(li.get_text(), style="ListNumber")
                    if color:
                        p.runs[0].font.color.rgb = RGBColor(*color)
            elif element.name == "a":  # Hyperlinks
                link_text = element.get_text()
                link_href = element.get("href")
                if link_text and link_href:
                    p = doc.add_paragraph()
                    add_hyperlink(p, link_text, link_href)
            elif element.name == "span":  # Colored text inside span
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                if color:
                    run.font.color.rgb = RGBColor(*color)
            elif element.name == "img":  # Images
                img_src = element.get("src")
                if img_src:
                    img_path = os.path.join(IMAGE_FOLDER, os.path.basename(img_src))
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(3))  # Resize image if needed
                        doc.add_paragraph(f"Image: {os.path.basename(img_src)}")

def generate_docx(folder_path, output_file):
    """Creates a Word document with the content extracted from HTML while maintaining tab structure."""
    if not os.path.exists(folder_path):
        print("Folder does not exist:", folder_path)
        return

    doc = Document()
    
    # Read frmPAVIS.html as the first page content
    index_file_path = os.path.join(folder_path, "frmPAVIS.html")
    if os.path.exists(index_file_path):
        print("Adding frmPAVIS.html as the first page...")
        extract_content_from_html(index_file_path, doc)
    else:
        print("Warning: Index file (frmPAVIS.html) not found!")

    # Process each HTML file (except frmPAVIS.html) while maintaining order
    for filename in sorted(os.listdir(folder_path)):  # Ensures order
        if filename.endswith(".html") and filename != "frmPAVIS.html":
            file_path = os.path.join(folder_path, filename)
            file_name = filename.replace(".html", "")

            print(f"Processing: {filename}")

            doc.add_page_break()
            doc.add_heading(file_name, level=1)
            extract_content_from_html(file_path, doc)

    # Save the Word document
    doc.save(output_file)
    print(f"Documentation created: {output_file}")

# Set folder path and output file
folder_path = r"C:\Users\somnath.das\Desktop\Default"  # Replace with your actual folder path
output_file = os.path.join(folder_path, "Application_Documentation.docx")

# Generate the document
generate_docx(folder_path, output_file)
